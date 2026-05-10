"""Generate DCW SIOP technical manual as Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).parent / "DCW_SIOP_Technical_Manual.docx"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h


def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def BULLET(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def NUM(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# ---- TITLE ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("DCW SIOP")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Standard-First Inventory & Order Planner")
sr.font.size = Pt(14)
sr.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run("Technical Manual").font.size = Pt(13)

doc.add_paragraph()

# ---- 1 OVERVIEW ----
H("1. Project Overview", 1)
P(
    "DCW SIOP is a standard-first inventory and order-planning tool for a pigment "
    "production line. The factory operator selects the standard currently running on "
    "the line; the system then routes open customer requirement lines for that color, "
    "ranks the inventory lots that can satisfy them using three independent matching "
    "models, and recommends an allocation against live stock before anything is "
    "committed."
)

H("1.1 Problem Solved", 2)
BULLET([
    "Operators previously chose lots manually; wrong picks shipped because two lots under the same standard can have very different color deltas.",
    "Demand was not routed by what was actually being produced, so allocations ignored production reality.",
    "Inventory shortage was discovered late, after commitment, instead of at planning time.",
    "There was no audit trail explaining why a specific lot was chosen over another.",
])

H("1.2 What the Project Achieves", 2)
BULLET([
    "Single operator input: pick the standard currently in production. Everything downstream derives from that choice.",
    "Color-routed demand: open requirement lines for that standard's color family load automatically.",
    "Model-driven lot ranking: every candidate lot is scored by Euclidean Delta-E, cosine similarity, and KNN distance, then merged into a consensus rank.",
    "Live stock simulation: lines are processed in fulfillability order against one shared live inventory pool, so the dashboard shows full, partial, and blocked outcomes before commit.",
    "Auditable recommendations: the Flow Audit page replays the engine decisions step by step, including the math for each top candidate.",
    "Commit lock-in: approved allocations decrement inventory, update line outstanding quantity, and append a commit audit row.",
    "Admin tooling: re-parse the inventory workbook, update the current production standard per color, and ingest invoices via OCR or manual entry.",
])

# ---- 2 ARCHITECTURE ----
H("2. System Architecture", 1)
P(
    "Two-tier app. Flask backend serves a JSON API; React (Create React App) frontend "
    "renders the dashboard. Identity is carried per request via the username header / "
    "query parameter so each operator sees their own session state."
)

TABLE(
    ["Layer", "Technology", "Role"],
    [
        ["Frontend", "React (CRA), JSX, plain CSS", "Dashboard, standard selector, results tabs, Flow Audit page"],
        ["Backend", "Flask, Flask-CORS", "JSON API, matching engine, allocation, admin endpoints"],
        ["Numerics", "NumPy, pandas, scikit-learn (NearestNeighbors, StandardScaler)", "Vector math, ranking, KNN"],
        ["Data store", "Stitched CSVs under backend/stitched_outputs/", "Inventory lots, lot test results, standard profiles, invoices"],
        ["Source workbook", "Main inventory data.xlsx", "Client-maintained inventory; parsed on startup or admin refresh"],
        ["OCR (optional)", "pytesseract, pdfplumber, pdf2image, OpenAI", "Invoice draft extraction"],
    ],
)

H("2.1 Repository Layout", 2)
CODE(
    "DCW SIOP/\n"
    "  backend/\n"
    "    app.py                  Flask app, matching engine, endpoints\n"
    "    Main inventory data.xlsx  Source inventory workbook\n"
    "    stitched_outputs/       Parsed CSVs (runtime data)\n"
    "  frontend/\n"
    "    src/\n"
    "      Dashboard.jsx         Top-level dashboard\n"
    "      components/\n"
    "        PigmentSelector.jsx   Standard selector\n"
    "        ProductionPanel.jsx   Inventory coverage panel\n"
    "        ResultsTabs.jsx       Lines / candidates / allocation / inventory tabs\n"
    "        MatchingFlowTab.jsx   Flow Audit page (engine walkthrough)\n"
    "        Sidebar.jsx           Standards, lots, lines, dataset mode\n"
    "        UpdateDataTab.jsx     Admin data update tab\n"
)

# ---- 3 DATA MODEL ----
H("3. Data Model", 1)
P("Stitched CSVs loaded on startup into in-memory pandas frames.")

TABLE(
    ["Dataset", "Key columns", "Purpose"],
    [
        ["inventory_lots", "lot_id, lot_no, grade, standard_code, qty_mt_on_hand, color_family", "Stock universe"],
        ["standard_profiles", "standard_code, grade, method_id, reference_l/a/b", "One current standard per color"],
        ["lot_test_results", "lot_no, standard_code, method_id, delta_l/a/b, delta_e, strength, absolute_l/a/b", "Per-lot QC test vectors"],
        ["invoice_headers", "invoice_id, invoice_number, customer_name", "Customer requirement headers"],
        ["invoice_lines", "invoice_line_id, invoice_id, grade, standard_code, application, qty_mt, target_method_id, target_delta_l/a/b", "Line-level demand and color target"],
    ],
)

# ---- 4 BACKEND ENDPOINTS ----
H("4. Backend API", 1)
TABLE(
    ["Method", "Endpoint", "Purpose"],
    [
        ["GET", "/api/standards", "Standards dropdown with color, inventory qty, lot count, methods"],
        ["GET", "/api/invoices", "Invoice coverage summary; supported = invoice color has stock"],
        ["POST", "/api/analyze/standard", "Core analysis. Returns eligible lines, candidate lots, allocation, inventory analysis"],
        ["POST", "/api/invoices/commit", "Commit allocation: decrement inventory, update outstanding qty, append audit"],
        ["GET/POST", "/api/admin/current-standards", "Read/update current production standard per color"],
        ["POST", "/api/admin/inventory/refresh", "Re-parse Main inventory data.xlsx into runtime CSVs"],
        ["POST", "/api/admin/invoice/ocr", "OCR/agent extracts invoice draft (header + lines + dL/dA/dB)"],
        ["POST", "/api/admin/invoice/manual", "Save reviewed invoice lines"],
    ],
)

# ---- 5 FLOW ----
H("5. Matching Flow (from Flow Audit page)", 1)
P(
    "The Flow Audit tab in MatchingFlowTab.jsx is the canonical walkthrough. Six "
    "numbered steps grouped into three phases: Scope, Rank, Outcome."
)

H("5.1 Phase: Scope", 2)
P("Step 1 - Choose the current standard.", bold=True)
P(
    "Operator selects the running standard. That choice defines which inventory lots "
    "and QC tests are available to the match engine. Requirement demand is then routed "
    "by the standard color family, not by a separate customer filter."
)
P("Step 2 - Load only open lines with remaining qty.", bold=True)
P(
    "Queue contains open requirement lines for that color only. Carried partials stay "
    "in the queue with their remaining quantity."
)

H("5.2 Phase: Rank", 2)
P("Step 3 - Resolve test method and target vector.", bold=True)
P(
    "Requirement-specific tests win when present. Otherwise matching starts from the "
    "base test and falls through only when availability requires it. dL / dA / dB from "
    "the requirement remain the primary scoring target for every candidate lot."
)
P("Step 4 - Score and filter candidate lots.", bold=True)
P(
    "Euclidean Delta-E, cosine similarity, and KNN distance are computed per lot, "
    "ranked into a consensus order, then filtered by the active tolerance band. The "
    "page audits both the starting eligible quantity and the live eligible quantity "
    "after earlier lines consume stock."
)

H("5.3 Phase: Outcome", 2)
P("Step 5 - Allocate against the shared live stock pool.", bold=True)
P(
    "Lines are processed in fulfillability order. Earlier lines can drain the best "
    "lots, which is how later lines become partial or fully blocked. Queue order is: "
    "carried partials first, then fulfillability rank, then line id."
)
P("Step 6 - Apply decisions, then lock in.", bold=True)
P(
    "The dashboard turns engine audit into operator actions: fulfill, partially "
    "fulfill, skip, or cannot fulfill. Lock-in decrements inventory and leaves "
    "shortfalls open. The Flow Audit page explains why the recommendation landed where "
    "it did before anything is committed."
)

# ---- 6 MATCHING MATH ----
H("6. Matching Models", 1)
P(
    "Three independent rankings per candidate lot, then a consensus rank. All three "
    "operate on the dL / dA / dB vector space."
)

H("6.1 Euclidean Delta-E", 2)
CODE("dE = sqrt( (lot.dL - tgt.dL)^2 + (lot.dA - tgt.dA)^2 + (lot.dB - tgt.dB)^2 )")
P("Smaller is better. Direct color distance.")

H("6.2 Cosine similarity", 2)
CODE("cos = (q . l) / (||q|| * ||l||)")
P("Direction match between target and lot delta vectors. Higher is better.")

H("6.3 KNN distance", 2)
P(
    "scikit-learn NearestNeighbors over the standardized dL/dA/dB feature space "
    "(StandardScaler). Distance from the target vector to each lot vector after "
    "scaling. Smaller is better."
)
CODE("knn = sqrt( sum( (lot_scaled[i] - tgt_scaled[i])^2 ) )")

H("6.4 Consensus rank", 2)
P(
    "Each lot gets euclideanRank, cosineRank, knnRank. Combined into a consensus "
    "score / consensusRank used as the primary lot ordering for allocation. Tie-break "
    "by available qty and stable lot number."
)

H("6.5 Tolerance bands", 2)
TABLE(
    ["Mode", "Acceptance"],
    [["strict", "Delta-E <= 1.5"], ["relaxed", "Delta-E <= 3.5"], ["review", "All scored candidates"]],
)

H("6.6 Super-lot guard", 2)
P(
    "A lot with more than two available test methods is tagged as a super lot. Super "
    "lots are reserved unless the requirement names every test on that lot, or "
    "non-super eligible stock cannot fulfill the line. Prevents over-consuming "
    "versatile inventory."
)

# ---- 7 ALLOCATION ----
H("7. Allocation", 1)
P(
    "Greedy allocation against one shared live inventory pool. Order: perceptual band, "
    "then consensus rank, then remaining qty, then lot number. Three outcomes possible "
    "per requirement line:"
)
TABLE(
    ["Outcome", "Meaning"],
    [
        ["Full", "Allocated qty == requested qty"],
        ["Partial", "Allocated qty > 0 and < requested qty (live pool drained earlier)"],
        ["Can't fulfill", "Allocated qty == 0 (stock shortage at this queue position)"],
        ["Unsupported", "Color not in current inventory scope at all"],
    ],
)

# ---- 8 FLOW AUDIT UI ----
H("8. Flow Audit Page (MatchingFlowTab.jsx)", 1)
P("The audit page renders six artifacts after an analysis runs:")
NUM([
    "Hero stats: color, open lines, full / partial / blocked counts, stock left.",
    "Six-step flow strip: scope, rank, outcome.",
    "Example A - full fill on first pass: shows target vector, consensus breakdown, math (Euclidean / Cosine / KNN), top-5 candidate table.",
    "Example B - partial because earlier lines consumed stock: shows starting vs live eligible MT, depleted lots, actual draws.",
    "Example C - blocked by inventory shortage: shows zero live eligible MT and empty allocation list.",
    "Queue audit table: every supported line in fulfillability order with starting elig, live elig, alloc MT, shortfall, outcome.",
])
P(
    "Guardrails section reminds the operator: shortages are inventory-backed, not "
    "synthetic; requirement rows preserve customer-facing grade/standard labels; one "
    "selected standard means one color pool and one shared stock draw."
)

# ---- 9 RUN ----
H("9. Running the App Locally", 1)
H("9.1 Backend", 2)
CODE(
    "cd backend\n"
    "pip install -r requirements.txt\n"
    "python app.py    # Flask on :5000\n"
)
H("9.2 Frontend", 2)
CODE(
    "cd frontend\n"
    "npm install\n"
    "npm start        # CRA dev server on :3000\n"
)
H("9.3 Default credentials", 2)
TABLE(
    ["Username", "Type"],
    [["Akash / Anirudh / Sanjay / Sushant / Naina", "user"], ["admin", "admin"]],
)

# ---- 10 ADMIN ----
H("10. Admin Workflows", 1)
BULLET([
    "Update current production standard: POST /api/admin/current-standards rewrites standard_profiles.csv and re-parses inventory.",
    "Refresh inventory: POST /api/admin/inventory/refresh re-parses Main inventory data.xlsx into inventory_lots.csv and lot_test_results.csv.",
    "Ingest invoice via OCR: POST /api/admin/invoice/ocr returns a draft header + lines (incl. dL/dA/dB) for review.",
    "Save reviewed invoice: POST /api/admin/invoice/manual appends to invoice_headers / invoice_lines.",
])

# ---- 11 ACHIEVEMENTS ----
H("11. Outcomes Achieved", 1)
BULLET([
    "Eliminated manual lot picking; operator workload reduced to a single dropdown choice.",
    "Replaced single-sort lot ordering with a three-model consensus, surfacing the correct lot when the standard alone is ambiguous.",
    "Made shortages visible at planning time via live-pool simulation, so partial / blocked lines are seen before commit.",
    "Produced a transparent audit (Flow Audit page) so allocation decisions are defensible to QC and customers.",
    "Wired admin tooling end-to-end: standard updates, inventory refresh, OCR + manual invoice ingest.",
    "Preserved customer-facing grade and standard labels on requirement rows while routing internally by color family.",
])

# ---- 12 GLOSSARY ----
H("12. Glossary", 1)
TABLE(
    ["Term", "Meaning"],
    [
        ["Standard", "Production color spec currently running on the line"],
        ["Color family", "Color group used to route requirement lines"],
        ["Lot", "Discrete inventory batch with QC test results"],
        ["dL / dA / dB", "Color delta components vs reference profile"],
        ["Delta-E", "Euclidean distance in dL/dA/dB space"],
        ["Consensus rank", "Combined rank from Euclidean, cosine, KNN"],
        ["Super lot", "Lot with > 2 available test methods; reserved by guard"],
        ["Live eligible MT", "Stock left in the pool when this line is processed"],
        ["Fulfillability rank", "Order in which requirement lines are processed"],
    ],
)

doc.save(OUT)
print(f"Wrote {OUT}")
